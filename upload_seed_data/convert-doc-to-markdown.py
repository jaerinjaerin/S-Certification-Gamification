import os
import json
import re
from docx import Document

def convert_bullets_to_html(cell_content):
    """
    Converts inline Markdown-style bullets in a table cell to HTML list items.
    """
    bullet_pattern = r'(\*|-)\s([^*|-]+)'  # Matches "* item" or "- item"
    matches = re.findall(bullet_pattern, cell_content)
    if matches:
        html_list = "".join([f"<li>{match[1].strip()}</li>" for match in matches])
        return f"<ul>{html_list}</ul>"
    return cell_content  # Return unchanged if no bullets found

def convert_markdown_table(markdown_content):
    """
    Processes a Markdown table and converts bullets in cells to HTML.
    """
    lines = markdown_content.split("\n")
    processed_lines = []
    for line in lines:
        if "|" in line:  # Table row
            cells = line.split("|")
            processed_cells = [convert_bullets_to_html(cell.strip()) for cell in cells]
            processed_lines.append(" | ".join(processed_cells))
        else:
            processed_lines.append(line)
    return "\n".join(processed_lines)

def convert_docx_to_markdown(doc_filepath):
    """
    Converts a .docx file to a Markdown string.
    """
    doc = Document(doc_filepath)
    markdown_content = ""

    for paragraph in doc.paragraphs:
        markdown_content += paragraph.text + "\n\n"  # Add paragraph text with double line breaks

    for table in doc.tables:
        markdown_content += "\n"
        for row in table.rows:
            row_data = " | ".join(cell.text.strip() for cell in row.cells)
            markdown_content += f"| {row_data} |\n"
        markdown_content += "\n"

    return markdown_content

def convert_excel_to_json_name(excel_filename):
    """
    Convert an Excel file name to a JSON file name.
    Example: "{Switzerland}.{NAT_2756}.{de-DE}.T&C.docx" -> "NAT_2756_de-DE.json"
    
    Args:
        excel_filename (str): The name of the Excel file.
        
    Returns:
        str: The converted JSON file name.
    """
    try:
        # Remove curly braces and split by dots
        cleaned_filename = re.sub(r"[{}]", "", excel_filename)
        parts = cleaned_filename.split('.')
        
        # Extract the second and third parts (e.g., NAT_2756 and de-DE)
        if len(parts) >= 3:
            json_name = f"{parts[1]}_{parts[2]}.json"
            return json_name
        else:
            raise ValueError("Invalid file name format")
    except Exception as e:
        return f"Error: {str(e)}"

def convert_doc_to_md_and_json(doc_filepath, markdown_dir, json_dir):
    """
    Converts a .docx file to Markdown and JSON formats with custom table processing.
    """
    # Convert .docx to Markdown
    markdown_content = convert_docx_to_markdown(doc_filepath)

    # Process table formatting
    markdown_content = convert_markdown_table(markdown_content)

    # Save Markdown content
    markdown_filename = os.path.splitext(os.path.basename(doc_filepath))[0] + ".md"
    markdown_filepath = os.path.join(markdown_dir, markdown_filename)
    with open(markdown_filepath, "w", encoding="utf8") as md_file:
        md_file.write(markdown_content)

    # Convert to JSON
    json_filename = convert_excel_to_json_name(os.path.basename(doc_filepath))
    json_filepath = os.path.join(json_dir, json_filename)

    json_content = json.dumps({"contents": markdown_content}, indent=4)

    # Save JSON content
    with open(json_filepath, "w", encoding="utf8") as json_file:
        json_file.write(json_content)

# Input directory
input_dir = "data/term-and-privacy/term/origins"

# Output directories
markdown_dir = "data/term-and-privacy/term/markdowns"
json_dir = "data/term-and-privacy/term/jsons"
os.makedirs(markdown_dir, exist_ok=True)
os.makedirs(json_dir, exist_ok=True)

for filename in os.listdir(input_dir):
    if filename.endswith(".doc") or filename.endswith(".docx"):
        input_filepath = os.path.join(input_dir, filename)
        print(f"Start converting {filename}")
        convert_doc_to_md_and_json(input_filepath, markdown_dir, json_dir)
        print(f"End converting {filename}")
